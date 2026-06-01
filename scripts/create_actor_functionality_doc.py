from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "DRRCS_Functionality_By_Actor.docx"


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document, items):
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_actor(document, actor, purpose, functions, workflow, data_access):
    document.add_heading(actor, level=1)
    document.add_paragraph(purpose)

    document.add_heading("Main Functionalities", level=2)
    add_bullets(document, functions)

    document.add_heading("Typical Workflow", level=2)
    add_numbered(document, workflow)

    document.add_heading("Data Access", level=2)
    add_bullets(document, data_access)


def main():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 2"].font.name = "Arial"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Disaster Relief & Resource Coordination System (DRRCS)")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Actor-Wise Functionality Document").bold = True

    document.add_paragraph(
        "This document explains the working and functionality of the DRRCS web application based on each system actor. "
        "The system supports disaster reporting, emergency help requests, resource management, shelter coordination, "
        "volunteer tracking, notifications, analytics, and secure role-based access."
    )

    document.add_heading("System Overview", level=1)
    add_bullets(
        document,
        [
            "Frontend: Next.js dashboard with login, registration, maps, analytics, and management sections.",
            "Backend: Node.js and Express REST API with JWT authentication and role-based access control.",
            "Database: PostgreSQL managed through Prisma ORM.",
            "Deployment: Docker Compose runs frontend, backend, and PostgreSQL services.",
            "Current demo location: Jalandhar, Punjab, with map fallback set to NIT Jalandhar.",
        ],
    )

    document.add_heading("Actors in the System", level=1)
    add_bullets(
        document,
        [
            "Admin",
            "Authority",
            "NGO Coordinator",
            "Volunteer",
            "Affected Individual",
        ],
    )

    add_actor(
        document,
        "1. Admin",
        "The Admin is the highest-level actor responsible for supervising the complete disaster relief platform.",
        [
            "Can log in securely using JWT authentication.",
            "Can view the main analytics dashboard.",
            "Can monitor active disasters, emergency requests, shelter occupancy, resources, and volunteers.",
            "Can manage disaster records such as floods and earthquakes.",
            "Can view and manage emergency help requests.",
            "Can manage resources such as food, water, medicine, blankets, and rescue equipment.",
            "Can view shelters and occupancy data.",
            "Can view registered volunteers and their availability.",
            "Can access exported CSV/PDF reports.",
            "Can view audit logs for important system activity.",
        ],
        [
            "Admin logs into the system.",
            "Admin opens the dashboard overview to inspect overall disaster status.",
            "Admin reviews active disasters and emergency requests.",
            "Admin allocates resources or coordinates with authorities/NGOs.",
            "Admin monitors volunteers and shelters until requests are resolved.",
        ],
        [
            "Full visibility over users, disasters, requests, resources, shelters, volunteers, notifications, and audit logs.",
            "Can perform high-level administrative actions such as deletion and system monitoring.",
        ],
    )

    add_actor(
        document,
        "2. Authority",
        "The Authority actor represents government or disaster-management officials who verify incidents and coordinate official response.",
        [
            "Can create and update disaster records.",
            "Can verify and update emergency request status.",
            "Can allocate resources based on priority.",
            "Can monitor shelter capacity and occupancy.",
            "Can review analytics related to active disasters and requests.",
            "Can send or participate in emergency broadcast communication.",
        ],
        [
            "Authority logs into the system.",
            "Authority checks reported disasters and help requests.",
            "Authority verifies whether requests are genuine and urgent.",
            "Authority updates request status from Pending to Assigned or In Progress.",
            "Authority coordinates resources, shelters, and volunteers.",
        ],
        [
            "Can access disaster, request, shelter, resource, and analytics data.",
            "Has operational access, but not unrestricted administrative control over all system records.",
        ],
    )

    add_actor(
        document,
        "3. NGO Coordinator",
        "The NGO Coordinator manages NGO-side relief contributions, resources, and volunteer coordination.",
        [
            "Can register and log in as an NGO Coordinator.",
            "Can add and track resources contributed by NGOs.",
            "Can monitor resource availability and distribution.",
            "Can view volunteer data for coordination.",
            "Can participate in task and relief operation planning.",
            "Can view dashboard analytics relevant to response operations.",
        ],
        [
            "NGO Coordinator logs into the system.",
            "Coordinator checks active disasters and resource shortages.",
            "Coordinator adds available NGO resources such as water, food, or medicine.",
            "Coordinator coordinates with volunteers and authorities.",
            "Coordinator tracks relief distribution until completion.",
        ],
        [
            "Can access resource, volunteer, disaster, request, and analytics information needed for NGO operations.",
            "Cannot perform full administrative actions reserved for Admin.",
        ],
    )

    add_actor(
        document,
        "4. Volunteer",
        "The Volunteer actor assists in field response, relief delivery, evacuation, and request handling.",
        [
            "Can register as a Volunteer.",
            "When a user registers as Volunteer, the system automatically creates a linked volunteer profile.",
            "Can log in and view their own volunteer record in the Volunteers section.",
            "Can view assigned tasks when task assignment is used.",
            "Can update task status such as Accepted, In Progress, or Completed.",
            "Can support delivery of resources, rescue operations, and shelter assistance.",
        ],
        [
            "Volunteer registers or logs into the system.",
            "System creates or verifies the volunteer profile.",
            "Volunteer opens the dashboard and checks assigned work.",
            "Volunteer performs the task on ground.",
            "Volunteer updates the task status after progress or completion.",
        ],
        [
            "Can view their own volunteer profile.",
            "Can view assigned tasks and related disaster/request details.",
            "Does not have full administrative access to all users or system settings.",
        ],
    )

    add_actor(
        document,
        "5. Affected Individual",
        "The Affected Individual represents a citizen or victim needing help during a disaster.",
        [
            "Can register and log in as an affected individual.",
            "Can view active disasters in the current region.",
            "Can submit emergency resource requests.",
            "Can request Food, Water, Medicine, Blankets, Shelter, Rescue Equipment, Evacuation Support, or Medical Assistance.",
            "Can set request priority as Low, Medium, High, or Urgent.",
            "Can describe the emergency need and submit it to authorities.",
            "Can track submitted request status.",
            "Can view shelters and resources available in the Jalandhar region.",
        ],
        [
            "Affected individual logs into the system.",
            "User opens the Requests section.",
            "User selects required resource type and priority.",
            "User enters details such as number of people affected and nearby landmark.",
            "System saves the request with Pending status.",
            "Authority/Admin reviews and processes the request.",
        ],
        [
            "Can see their own emergency requests.",
            "Can view public disaster, shelter, and resource information.",
            "Cannot modify other users, resources, disasters, or volunteer assignments.",
        ],
    )

    document.add_heading("Important Modules", level=1)
    modules = {
        "Authentication Module": [
            "Supports register, login, logout, forgot password, reset password, and current-user lookup.",
            "Uses bcrypt password hashing and JWT tokens.",
            "Stores token in the frontend and verifies it through the backend.",
        ],
        "Disaster Management Module": [
            "Stores disaster title, type, description, location, latitude, longitude, severity, start date, and status.",
            "Current sample disasters are Jalandhar Urban Flood Alert and Minor Earthquake Near NIT Jalandhar.",
        ],
        "Emergency Request Module": [
            "Allows affected individuals to submit help/resource requests.",
            "Supports request status flow: Pending, Assigned, In Progress, Resolved.",
        ],
        "Resource Management Module": [
            "Tracks relief resources such as water, medicine, blankets, food, and rescue equipment.",
            "Current resource locations are set around Jalandhar and NIT Jalandhar.",
        ],
        "Shelter Management Module": [
            "Tracks shelter name, address, capacity, occupied beds, GPS coordinates, and contact information.",
            "Current shelters are based in Jalandhar, Punjab.",
        ],
        "GIS/Map Module": [
            "Shows disaster, shelter, request, and location markers.",
            "Uses current browser location if allowed; otherwise falls back to NIT Jalandhar.",
        ],
        "Analytics and Reporting Module": [
            "Shows active disasters, available resources, shelter occupancy, open requests, and volunteer statistics.",
            "Supports CSV and PDF report export.",
        ],
    }
    for module, points in modules.items():
        document.add_heading(module, level=2)
        add_bullets(document, points)

    document.add_heading("Database Usage", level=1)
    add_bullets(
        document,
        [
            "PostgreSQL stores all persistent application data.",
            "Prisma ORM maps application models to database tables.",
            "Main tables include User, Role, Disaster, EmergencyRequest, Resource, Shelter, Volunteer, Task, Notification, NGO, and AuditLog.",
            "Docker volume storage keeps the database saved even after stopping and restarting containers.",
            "Data is lost only if Docker volumes are deleted using commands such as docker compose down -v.",
        ],
    )

    document.add_heading("Conclusion", level=1)
    document.add_paragraph(
        "DRRCS provides a role-based disaster relief coordination workflow. Each actor has clear responsibilities: "
        "affected individuals request help, volunteers assist on ground, NGOs provide resources, authorities verify "
        "and coordinate response, and admins supervise the complete operation. The system is suitable for demonstrating "
        "a final-year software engineering project with authentication, database persistence, maps, dashboards, and Docker deployment."
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
